from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RESEARCH = ROOT / "docs" / "current-research"
OUT_DIR = ROOT / "Presentation Materials"
OUT_DIR.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(0x14, 0x2A, 0x4A)
ORANGE = RGBColor(0xE8, 0x7A, 0x2E)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def title_page(doc: Document, title: str, subtitle: str, meta_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = ORANGE

    for line in meta_lines:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(line)
        r3.font.size = Pt(10.5)
    doc.add_paragraph()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def note_box(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = ORANGE
    p.add_run(text).italic = True


def equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(12)
    r.bold = True


def add_image(doc: Document, path: Path, caption: str, width_in: float = 6.3) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9.5)
    else:
        para(doc, f"[Figure missing: {path.name}] {caption}")


def make_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# DOCUMENT 1 — Atomic-level architecture report
# ---------------------------------------------------------------------------

def build_architecture_doc() -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    title_page(
        doc,
        "MatriWatch — Complete Atomic-Level Architecture Reference",
        "Dual-Scale Fine-Grained Postpartum Depression Screening Using a Score-Free Task-Adaptive Hybrid Framework",
        [
            "Authors: Arnob Aich Anurag, Shamiul Islam, Sadia Sultana — Team MatriWatch",
            "Event: Research Symposium 2.0, National Abstract Competition, AUST IEOM Student Chapter",
            "Document purpose: a complete, plain-language, every-detail technical reference of the model architecture",
        ],
    )
    page_break(doc)

    # 0. How to read this
    h1(doc, "0. How to Read This Document")
    para(
        doc,
        "This document is written so that every number, every layer, and every design choice in the MatriWatch "
        "postpartum depression (PPD) screening system is explained in plain language, with no detail skipped. "
        "Wherever a number comes directly from the saved model files in the project (the .pth checkpoints and "
        ".joblib model files), it is reported exactly as found. Wherever a value is a standard convention that is "
        "not literally stored in the saved file (for example, the exact dropout probability, since dropout layers "
        "carry no learnable weights), this is stated honestly as an assumption rather than presented as a measured fact.",
    )
    bullets(
        doc,
        [
            "Section 1–2: the problem, the goal, and the people behind the research.",
            "Section 3: the dataset, atom by atom — every one of the 46 input questions used by the model.",
            "Section 4: the exact preprocessing pipeline, step by step.",
            "Section 5: the Attention-Based 1D-CNN branch — every layer, every shape, every parameter count.",
            "Section 6: the Tuned RBF-SVM branch — every hyperparameter and support-vector count.",
            "Section 7: how the system decides which branch to trust (the Task-Adaptive Selector).",
            "Section 8–9: the final outputs, the measured accuracy, and what the figures show.",
            "Section 10: a plain-English glossary of every technical term used.",
        ],
    )

    # 1. Identity
    h1(doc, "1. Project Identity")
    make_table(
        doc,
        ["Field", "Value"],
        [
            ["Title", "Dual-Scale Fine-Grained Screening of Postpartum Depression Using a Score-Free Task-Adaptive Hybrid Framework in Bangladesh"],
            ["Team", "Team MatriWatch — Arnob Aich Anurag, Shamiul Islam, Sadia Sultana"],
            ["Keywords", "Postpartum Depression, PHQ-9, EPDS, 1D-CNN, Bangladesh"],
            ["Competition", "Research Symposium 2.0 — National Abstract Competition (AUST IEOM Student Chapter)"],
            ["Cohort", "800 postpartum women, Bangladesh-focused"],
            ["Two prediction targets", "PHQ-9 severity (6 classes) and EPDS risk level (3 classes: Low / Medium / High)"],
        ],
    )

    # 2. Problem & goal
    h1(doc, "2. The Clinical Problem and the Research Goal")
    para(
        doc,
        "Postpartum depression (PPD) affects mothers, infants, and family well-being, but in low-resource settings "
        "such as Bangladesh, early detection is still limited. Standard screening tools — the PHQ-9 (Patient Health "
        "Questionnaire-9) and the EPDS (Edinburgh Postnatal Depression Scale) — are widely used, but most prior "
        "machine-learning studies only do binary detection (depressed vs. not depressed) or simply re-predict the "
        "raw questionnaire total score. MatriWatch instead targets fine-grained, multi-class severity prediction "
        "for both scales at once, using a single shared set of psychosocial and clinical survey answers, while "
        "deliberately removing the raw PHQ-9 and EPDS numeric scores from the input features (a 'score-free' "
        "setting) so that the model must learn from the underlying life circumstances rather than from the "
        "questionnaire arithmetic itself.",
    )
    bullets(
        doc,
        [
            "Build a Bangladesh-focused postpartum depression screening model.",
            "Predict PHQ-9 severity (6 classes) and EPDS risk level (3 classes).",
            "Compare an Attention-based 1D-CNN branch against a tuned RBF-SVM branch.",
            "Handle missing values and class imbalance correctly.",
            "Explain every prediction using SHAP values and feature importance, so the model is not a black box.",
        ],
    )
    note_box(
        doc,
        "Research gap addressed",
        "Most existing work either performs binary PPD detection or predicts the raw screening score directly. "
        "Very little prior work attempts fine-grained, score-free, dual-scale (PHQ-9 + EPDS) classification on a "
        "Bangladesh-specific postpartum cohort with explainable machine learning.",
    )

    # 3. Dataset
    h1(doc, "3. The Dataset, Atom by Atom")
    para(
        doc,
        "The raw survey instrument collects 51 columns per respondent. Of these, 46 columns are predictor "
        "(input) questions, and the remaining 4 columns are the two raw screening outcomes: 'PHQ9 Score' + "
        "'PHQ9 Result' and 'EPDS Score' + 'EPDS Result'. In the score-free design used by the final pipeline, "
        "the two raw numeric scores (PHQ9 Score and EPDS Score) are excluded from the model's input features — "
        "only the 46 underlying survey questions are used to predict the severity/result labels. This is why the "
        "model input width is exactly 46.",
    )
    h2(doc, "3.1 The Full List of 46 Input Features")
    para(doc, "Grouped here by domain for readability; the model itself treats them as one flat 46-length vector.")

    feature_groups = {
        "Demographic & socioeconomic (10)": [
            "Age", "Current residence", "Level of education", "Marital status",
            "Occupation before latest pregnancy", "Monthly income before latest pregnancy",
            "Occupation after latest childbirth", "Current monthly income",
            "Husband's level of education", "Husband's monthly income",
        ],
        "Household & relational (8)": [
            "Addiction (any form)", "Number of total children", "Family type",
            "Number of household members", "Relationship with the in-laws",
            "Relationship with husband", "Relationship/bonding with the newborn",
            "Relationship between husband and child",
        ],
        "Psychosocial & support (6)": [
            "Feeling about motherhood", "Support received before/after childbirth",
            "Need for more support", "Major changes or losses during pregnancy",
            "Experience of abuse/mistreatment", "Trust and sharing feelings with close friends",
        ],
        "Obstetric history (5)": [
            "History of pregnancy loss", "Number of the latest pregnancy", "Pregnancy length (weeks)",
            "Was this pregnancy planned", "Had regular checkups",
        ],
        "Clinical / medical (4)": [
            "Disease before pregnancy", "Fear of pregnancy", "Diseases during pregnancy",
            "Mode of delivery",
        ],
        "Newborn & postpartum (10)": [
            "Age of newborn", "Age of older children (if any)", "Gender of newborn",
            "Birth complications", "Breastfeeding status", "Newborn illness",
            "Worry about newborn's health", "Can relax/sleep when newborn is tended by someone else",
            "Can relax/sleep when newborn is asleep", "Anger/irritation after childbirth, difficult to calm down",
        ],
        "Functioning & prior mental health (3)": [
            "Feeling about resuming regular activities", "Depression before pregnancy (PHQ-2)",
            "Depression during pregnancy (PHQ-2)",
        ],
    }
    for group, items in feature_groups.items():
        h3(doc, group)
        bullets(doc, items)
    total_count = sum(len(v) for v in feature_groups.values())
    note_box(doc, "Count check", f"{total_count} features grouped above = 46 total model inputs.")

    h2(doc, "3.2 The Two Prediction Targets (Excluded from Input)")
    make_table(
        doc,
        ["Target", "Classes", "Source column", "Excluded raw score"],
        [
            ["PHQ-9 severity", "6 classes (fine-grained PHQ-9 severity bands)", "PHQ9 Result", "PHQ9 Score (excluded — score-free)"],
            ["EPDS risk level", "3 classes: Low / Medium / High", "EPDS Result", "EPDS Score (excluded — score-free)"],
        ],
    )
    note_box(
        doc,
        "Honesty note",
        "The exact text label assigned to each of the 6 PHQ-9 output indices (0–5) is set inside the original "
        "training notebook's label encoder and is not re-stored inside the saved .pth checkpoint files in this "
        "repository. The EPDS 3-class mapping (Low / Medium / High) is confirmed directly from the dataset's "
        "'EPDS Result' column and the poster's EPDS severity pie chart.",
    )

    h2(doc, "3.3 Class Balance and Splitting")
    add_image(doc, RESEARCH / "dataset_description.png", "Figure 1. Dataset composition, raw class distributions, and the resampled train/validation/test split.")
    add_image(doc, RESEARCH / "eda_overview.png", "Figure 2. Exploratory profile — PHQ-9 severity counts, EPDS proportions, and the PHQ-9 score histogram.")
    para(
        doc,
        "Out of 800 respondents, both PHQ-9 and EPDS labels are imbalanced (Moderate and Mild dominate PHQ-9; "
        "High is the largest EPDS group at 43.8%, Low 32.5%, Medium 23.8%). The data is split in a stratified "
        "70% / 15% / 15% ratio for train / validation / test — approximately 560 training rows, 120 validation "
        "rows, and 120 test rows before SMOTE rebalancing is applied to the training portion.",
    )

    # 4. Preprocessing
    h1(doc, "4. Preprocessing Pipeline — Step by Step")
    add_image(doc, RESEARCH / "methodology_pipeline.png", "Figure 3. The full score-free task-adaptive hybrid pipeline, end to end.")
    numbered(
        doc,
        [
            "Start from the raw survey table of 800 rows x 51 columns.",
            "Drop the respondent identifier and both raw total-score columns (PHQ9 Score, EPDS Score) — this is what makes the pipeline 'score-free'.",
            "Mode imputation: any missing categorical answer is filled with that column's most frequent (mode) value.",
            "Label encoding: every categorical/text answer (e.g. 'Married', 'Joint family', 'Yes/No') is converted to an integer code.",
            "Standardization: every numeric feature is rescaled to zero mean and unit variance.",
            "SMOTE (Synthetic Minority Over-sampling Technique) is applied to the training split only, to synthesize extra examples of the rarer severity classes so the model is not biased toward the majority class.",
            "Stratified split into train (70%) / validation (15%) / test (15%), preserving class proportions in each subset.",
            "The resulting 46-column numeric matrix is fed in parallel to the CNN branch (as a 1-channel sequence) and the SVM branch (as a flat feature vector).",
        ],
    )
    equation(doc, "Standardization:   z = (x − μ) / σ")
    para(doc, "where x is the raw feature value, μ is that feature's training-set mean, and σ is its training-set standard deviation.")

    # 5. CNN branch
    h1(doc, "5. Branch A — Attention-Based 1D-CNN (Layer by Layer, Atomic Detail)")
    para(
        doc,
        "This branch treats the 46 score-free features as a length-46, 1-channel sequence (think of it as a "
        "miniature 1-D 'signal' of 46 data points) and learns local patterns across neighboring features using "
        "1-D convolutions, then applies a lightweight channel-attention gate, then a fully-connected (MLP) "
        "classification head. The exact layer shapes below were extracted directly from the saved PyTorch "
        "checkpoint files (ppd_cnn_phq9_best.pth and ppd_cnn_epds_best.pth) — every number is real, not estimated.",
    )

    h2(doc, "5.1 Input")
    bullets(
        doc,
        [
            "Input tensor shape: (batch_size, 1 input channel, 46 sequence positions).",
            "Each of the 46 positions corresponds to exactly one preprocessed survey feature.",
        ],
    )

    h2(doc, "5.2 Convolutional Feature Extractor — 3 Blocks")
    para(doc, "Each block follows the pattern: Conv1d (no bias) -> BatchNorm1d -> ReLU activation.")
    make_table(
        doc,
        ["Block", "Layer", "Weight shape (out, in, kernel)", "Kernel size", "Parameters", "Followed by"],
        [
            ["conv1", "Conv1d(1 -> 32)", "(32, 1, 5)", "5", "160", "BatchNorm1d(32) + ReLU"],
            ["conv2", "Conv1d(32 -> 64)", "(64, 32, 5)", "5", "10,240", "BatchNorm1d(64) + ReLU"],
            ["conv3", "Conv1d(64 -> 128)", "(128, 64, 5)", "5", "40,960", "BatchNorm1d(128) + ReLU"],
        ],
    )
    note_box(
        doc,
        "Assumption flagged",
        "Padding is not stored inside the saved weight tensors (only kernel weights and BatchNorm statistics are "
        "saved). The architecture is consistent with 'same' padding (padding = 2 for a kernel of 5) so the sequence "
        "length of 46 is preserved through all three convolution blocks before pooling — this is the standard "
        "convention for this block pattern, but is stated here as an inference, not a literal stored value.",
    )

    h2(doc, "5.3 Squeeze-and-Excitation Style Channel Attention")
    para(
        doc,
        "After the third convolution block produces a (batch, 128 channels, 46 positions) feature map, the "
        "network performs Global Average Pooling across the 46 positions to get one 128-length summary vector "
        "per sample. This vector is then passed through a small 'gate' network that learns how much to trust "
        "each of the 128 channels:",
    )
    make_table(
        doc,
        ["Layer", "Weight shape", "Parameters", "Activation"],
        [
            ["attn.gate[0] — Linear(128 -> 16)", "(16, 128)", "2,048", "ReLU (squeeze)"],
            ["attn.gate[2] — Linear(16 -> 128)", "(128, 16)", "2,048", "Sigmoid (excite, gate weights in [0, 1])"],
        ],
    )
    equation(doc, "attention_weights = σ( W₂ · ReLU( W₁ · GAP(features) ) )")
    equation(doc, "recalibrated_features = features × attention_weights   (channel-wise multiply)")
    para(
        doc,
        "In plain language: the network squeezes the 128 channels down to 16 (a bottleneck), decides which "
        "channels matter most, expands back to 128 channel weights between 0 and 1 (sigmoid), and multiplies "
        "the original 128-length pooled feature vector by these weights. Channels that matter more for this "
        "particular patient's pattern get amplified; channels that matter less get suppressed. Neither linear "
        "layer in the gate uses a bias term, which is the standard lightweight Squeeze-Excitation design.",
    )

    h2(doc, "5.4 Classification Head (MLP) — 3 Hidden Blocks + Output Layer")
    para(
        doc,
        "The attention-recalibrated 128-length vector is fed into a 4-layer fully connected classifier. Every "
        "hidden block follows the same pattern: Linear -> BatchNorm1d -> ReLU -> Dropout. This pattern is "
        "confirmed structurally: there are learnable parameters only at indices 0, 1, 4, 5, 8, 9, and 12 of the "
        "classifier, with exactly two parameter-free layers (ReLU then Dropout) sitting between each pair of "
        "Linear+BatchNorm blocks — that gap pattern is the fingerprint of a Dropout layer following each ReLU.",
    )
    make_table(
        doc,
        ["Index", "Layer", "Shape", "Parameters", "Note"],
        [
            ["classifier.0", "Linear(128 -> 256)", "(256, 128) + bias 256", "33,024", "Hidden block 1"],
            ["classifier.1", "BatchNorm1d(256)", "256 scale + 256 shift", "512 trainable", "+ running stats (non-trainable)"],
            ["classifier.2", "ReLU", "—", "0", "Activation"],
            ["classifier.3", "Dropout", "—", "0", "Regularization (rate not stored in weights)"],
            ["classifier.4", "Linear(256 -> 128)", "(128, 256) + bias 128", "32,896", "Hidden block 2"],
            ["classifier.5", "BatchNorm1d(128)", "128 scale + 128 shift", "256 trainable", "+ running stats"],
            ["classifier.6", "ReLU", "—", "0", "Activation"],
            ["classifier.7", "Dropout", "—", "0", "Regularization"],
            ["classifier.8", "Linear(128 -> 64)", "(64, 128) + bias 64", "8,256", "Hidden block 3"],
            ["classifier.9", "BatchNorm1d(64)", "64 scale + 64 shift", "128 trainable", "+ running stats"],
            ["classifier.10", "ReLU", "—", "0", "Activation"],
            ["classifier.11", "Dropout", "—", "0", "Regularization"],
            ["classifier.12", "Linear(64 -> N classes)", "PHQ-9: (6,64)+6 / EPDS: (3,64)+3", "PHQ-9: 390 / EPDS: 195", "Output logits layer"],
        ],
    )
    note_box(
        doc,
        "Dropout rate honesty note",
        "Dropout layers carry no learnable weights, so the exact dropout probability (a common choice would be "
        "0.3–0.5) is a training-time hyperparameter that is not recoverable from the saved .pth file. What is "
        "confirmed with certainty from the checkpoint structure is that a Dropout layer exists after every ReLU "
        "in the classifier head (3 dropout layers total), which is a strong, standard regularization choice for "
        "a small tabular dataset of only ~560 effective training rows.",
    )

    h2(doc, "5.5 Output Layer")
    bullets(
        doc,
        [
            "PHQ-9 head: final Linear(64 -> 6), producing 6 raw class scores (logits), converted to probabilities by Softmax.",
            "EPDS head: final Linear(64 -> 3), producing 3 raw class scores (logits), converted to probabilities by Softmax.",
        ],
    )
    equation(doc, "P(class = k | x) = exp(logit_k) / Σⱼ exp(logit_j)")

    h2(doc, "5.6 Total Parameter Count (Directly Counted From the Saved Checkpoints)")
    make_table(
        doc,
        ["Branch", "Trainable parameters", "Non-trainable BatchNorm running-statistic buffers", "Total stored tensors"],
        [
            ["PHQ-9 CNN (ppd_cnn_phq9_best.pth)", "131,366", "1,350", "132,716"],
            ["EPDS CNN (ppd_cnn_epds_best.pth)", "131,171", "1,350", "132,521"],
        ],
    )
    para(
        doc,
        "The two branches are almost identical in size — the only structural difference between them is the "
        "very last layer (6 output classes for PHQ-9 vs. 3 output classes for EPDS), a difference of exactly "
        "195 parameters, which matches precisely what is measured (132,716 − 132,521 = 195).",
    )

    h2(doc, "5.7 Training Behaviour (From the Results Report)")
    add_image(doc, RESEARCH / "training_curves_phq9.png", "Figure 4. PHQ-9 training curves — loss decreasing steadily, validation accuracy and weighted F1 rising and stabilizing.")
    add_image(doc, RESEARCH / "training_curves_epds.png", "Figure 5. EPDS training curves — similarly stable convergence for the 3-class task.")
    para(
        doc,
        "Both branches show smooth, non-oscillatory convergence with no signs of catastrophic overfitting — "
        "validation accuracy and weighted F1 track each other closely, which suggests the minority severity "
        "classes are being learned rather than ignored.",
    )

    # 6. SVM branch
    h1(doc, "6. Branch B — Tuned RBF-SVM (Exact Hyperparameters and Support Vectors)")
    para(
        doc,
        "The second branch is a Support Vector Machine with a Radial Basis Function (RBF) kernel — a classical, "
        "non-neural decision boundary learner that, in this project, repeatedly outperformed the CNN branch on "
        "validation weighted F1 for this tabular, relatively small dataset.",
    )
    equation(doc, "K(x, x′) = exp( −γ · ‖x − x′‖² )")
    para(doc, "γ (gamma) controls how far the influence of a single training example reaches; C controls how much the model is penalized for misclassifying a point (a tighter, more complex boundary vs. a smoother one).")

    h2(doc, "6.1 Score-Free Task-Adaptive SVM (46 input features — matches the poster's 'Tuned RBF-SVM Decision Branch')")
    make_table(
        doc,
        ["Task", "Kernel", "C", "gamma", "Input features", "Classes", "Total support vectors", "Support vectors per class"],
        [
            ["PHQ-9 (6-class)", "RBF", "2.0", "0.1", "46", "6", "814", "166, 136, 166, 164, 31, 151"],
            ["EPDS (3-class)", "RBF", "3.0", "0.05", "46", "3", "670", "245, 217, 208"],
        ],
    )

    h2(doc, "6.2 Enriched-Fusion SVM (49 input features — the higher-performing final configuration shown on the poster's comparison charts)")
    para(
        doc,
        "A later iteration uses 3 additional engineered/fused features on top of the original 46 (49 total "
        "inputs), paired with a much larger C (more confidence in fitting the training data tightly) and a "
        "smaller gamma (a smoother, wider-reaching decision surface). This is the version whose bar reaches the "
        "highest accuracy and weighted F1 on the poster's 'Enriched Fusion Final' comparison charts.",
    )
    make_table(
        doc,
        ["Task", "Kernel", "C", "gamma", "Input features", "Classes", "Total support vectors", "Support vectors per class"],
        [
            ["PHQ-9 (6-class)", "RBF", "20.0", "0.01", "49", "6", "622", "150, 89, 162, 133, 13, 75"],
            ["EPDS (3-class)", "RBF", "20.0", "0.005", "49", "3", "361", "85, 109, 167"],
        ],
    )
    note_box(
        doc,
        "What the 3 extra features likely are",
        "The exact identity of the 3 additional 'enriched fusion' features is not preserved as metadata inside "
        "the saved .joblib file, and the training script that produced this specific artifact is not present in "
        "this repository (only the resulting trained model file is). Based on the name 'enriched fusion' and the "
        "poster's framing, the most defensible description is that this configuration fuses a small number of "
        "additional engineered or CNN-derived signals on top of the original 46 score-free survey features — this "
        "should be confirmed against the original training notebook before being stated as definitive in a paper.",
    )
    bullets(
        doc,
        [
            "Both SVMs use probability=True, meaning Platt scaling is applied on top of the raw decision function so the model can output class probabilities, not just hard labels.",
            "random_state=42 is fixed for reproducibility in every variant.",
            "Support vectors are the actual training examples that sit on or inside the margin and therefore define the decision boundary — the SVM does not need to remember the other training rows once support vectors are identified.",
        ],
    )

    # 7. Task-adaptive selector
    h1(doc, "7. The Task-Adaptive Selector — How the System Chooses a Winner")
    para(
        doc,
        "Rather than always trusting one model, the pipeline trains the CNN branch and the SVM branch in "
        "parallel for each task, evaluates both on the held-out validation split using weighted F1-score, and "
        "deploys whichever branch scores higher — independently for PHQ-9 and for EPDS. This is why the method is "
        "called 'task-adaptive': the winning architecture is not fixed in advance, it is chosen per task based on "
        "evidence.",
    )
    numbered(
        doc,
        [
            "Train the Attention-CNN branch on the training split.",
            "Train the tuned RBF-SVM branch on the same training split.",
            "Evaluate both branches on the validation split using weighted F1-score.",
            "For PHQ-9: keep whichever branch has the higher validation weighted F1.",
            "For EPDS: repeat the same comparison independently.",
            "Report final test-set metrics only from the selected (winning) branch for each task.",
        ],
    )
    note_box(
        doc,
        "What actually won, per the Results Report",
        "In the reported score-free experiment run, the tuned RBF-SVM branch outperformed the Attention-CNN "
        "branch on validation for both PHQ-9 and EPDS, so the deployed predictions for both tasks come from the "
        "SVM branch (or its enriched-fusion successor) rather than from the CNN branch alone in that run.",
    )

    # 8. Outputs
    h1(doc, "8. Final Model Outputs")
    make_table(
        doc,
        ["Task", "Output classes", "Class meaning"],
        [
            ["PHQ-9", "6", "Fine-grained PHQ-9 severity bands (exact text labels are set by the original notebook's label encoder)"],
            ["EPDS", "3", "Low risk / Medium risk / High risk"],
        ],
    )

    # 9. Metrics
    h1(doc, "9. Measured Performance")
    h2(doc, "9.1 Final Test-Set Metrics (Score-Free Task-Adaptive Hybrid, as reported in the Results Report)")
    make_table(
        doc,
        ["Task", "Accuracy", "Weighted F1", "Weighted Precision", "Weighted Recall"],
        [
            ["PHQ-9 (6-class)", "80.37%", "81.61%", "85.56%", "80.37%"],
            ["EPDS (3-class)", "82.91%", "83.04%", "83.58%", "82.91%"],
        ],
    )
    h2(doc, "9.2 PHQ-9 Baseline Comparison (score-free setting)")
    make_table(
        doc,
        ["Model", "Accuracy", "Weighted F1"],
        [
            ["Logistic Regression", "58.88%", "—"],
            ["Random Forest", "70.09%", "—"],
            ["Gradient Boosting", "56.54%", "—"],
            ["Generic RBF-SVM (untuned)", "77.57%", "76.68%"],
            ["Attention-CNN branch alone", "67.76%", "68.88%"],
            ["Final selected task-adaptive hybrid", "80.37%", "81.61%"],
        ],
    )
    add_image(doc, RESEARCH / "baseline_comparison.png", "Figure 6. Baseline comparison chart for the PHQ-9 task (score-free setting).")
    note_box(
        doc,
        "Poster's further-improved figure",
        "The poster's 'Enriched Fusion Final' bars show an even higher accuracy and weighted F1 (visually "
        "reading approximately the low-90s percent for both metrics on PHQ-9), corresponding to the "
        "49-feature enriched-fusion SVM described in Section 6.2. This is the most recent and strongest "
        "iteration of the pipeline; precise decimal figures should be re-confirmed from the original notebook "
        "output before being quoted in a publication, since the poster chart is an image, not machine-readable text.",
    )

    h2(doc, "9.3 Confusion Matrices")
    add_image(doc, RESEARCH / "phq9_confusion_matrix.png", "Figure 7. PHQ-9 confusion matrix — strongly diagonal; remaining errors are mostly between clinically adjacent severity levels.")
    add_image(doc, RESEARCH / "epds_confusion_matrix.png", "Figure 8. EPDS confusion matrix — dominant diagonal; misclassifications stay near neighboring risk bands.")

    h2(doc, "9.4 ROC Curves")
    add_image(doc, RESEARCH / "roc_phq9.png", "Figure 9. One-vs-rest ROC curves for the 6 PHQ-9 classes.")
    add_image(doc, RESEARCH / "roc_epds.png", "Figure 10. One-vs-rest ROC curves for the 3 EPDS classes.")

    h2(doc, "9.5 Class Distributions")
    add_image(doc, RESEARCH / "donuts_phq9.png", "Figure 11. PHQ-9 class proportions.")
    add_image(doc, RESEARCH / "donuts_epds.png", "Figure 12. EPDS class proportions.")

    h1(doc, "10. Explainability — Why the Model Decides What It Decides")
    add_image(doc, RESEARCH / "feature_importance.png", "Figure 13. Random Forest feature importance, score-free PHQ-9 setting.")
    add_image(doc, RESEARCH / "shap_summary.png", "Figure 14. SHAP summary plot for multiclass PHQ-9 prediction.")
    para(
        doc,
        "In the score-free setting, the strongest predictors of PHQ-9 severity are not demographic facts like "
        "age alone, but relational and postpartum-context variables: relationship with the newborn, relationship "
        "with husband, relationship between father and newborn, anger after the latest childbirth, relationship "
        "with the in-laws, age of newborn, occupation, received support, and experience of abuse. SHAP confirms "
        "that these same variables consistently push predictions toward higher or lower severity across many "
        "patients, and that the direction of the effect (helpful vs. harmful) depends on the specific value each "
        "patient reports — this is what makes the model interpretable rather than a black box.",
    )

    h1(doc, "11. Plain-English Glossary")
    make_table(
        doc,
        ["Term", "Plain explanation"],
        [
            ["1D-CNN", "A neural network that slides small filters across a sequence of numbers to find local patterns, the way a 2D-CNN slides filters across an image."],
            ["BatchNorm1d", "Rescales the values flowing through a layer so they have a consistent mean/spread, which speeds up and stabilizes training."],
            ["Dropout", "Randomly switches off some neurons during training only, forcing the network not to over-rely on any single feature path; helps prevent overfitting."],
            ["Squeeze-and-Excitation (SE) attention", "A tiny side-network that looks at the overall pattern of channels and decides which ones to amplify or suppress for this specific input."],
            ["SMOTE", "Synthetic Minority Over-sampling Technique — creates new, realistic synthetic examples of rare classes so the model sees them often enough to learn them."],
            ["RBF-SVM", "A Support Vector Machine using a Radial-Basis-Function kernel, which can draw curved (non-linear) decision boundaries between classes."],
            ["Support vector", "A training example that lies close enough to the decision boundary that it directly shapes where that boundary is drawn."],
            ["Weighted F1-score", "A single accuracy-like score that balances precision and recall, weighted by how many examples each class actually has — fairer than raw accuracy on imbalanced data."],
            ["SHAP value", "A number that explains how much one specific feature pushed one specific prediction up or down, based on game-theory principles."],
            ["Score-free", "A design choice where the raw PHQ-9/EPDS numeric totals are deliberately excluded from the model's inputs, forcing it to learn from life-context answers instead of the questionnaire arithmetic."],
            ["Task-adaptive", "The deployed model is chosen per prediction task (PHQ-9 vs. EPDS) based on which branch performs better on validation data, rather than being fixed in advance."],
        ],
    )

    out_path = OUT_DIR / "MatriWatch_Atomic_Architecture_Detailed_Report.docx"
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# DOCUMENT 2 — Presentation speech
# ---------------------------------------------------------------------------

def build_speech_doc() -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    title_page(
        doc,
        "MatriWatch — Final-Round Presentation Speech",
        "Research Symposium 2.0 — National Abstract Competition, AUST IEOM Student Chapter",
        [
            "Team MatriWatch: Arnob Aich Anurag, Shamiul Islam, Sadia Sultana",
            "Format: 5-minute spoken presentation + judges' Q&A",
            "Judging criteria: Clarity & Organization, Research Quality & Innovation, Visual Appeal, Presentation Skills, Industry Relevance",
        ],
    )
    page_break(doc)

    h1(doc, "How to Use This Script")
    para(
        doc,
        "This speech is written to fit a strict 5-minute on-site poster presentation, followed by a judges' Q&A. "
        "It follows the exact section order suggested in the competition rulebook (Title, Literature Review, "
        "Research Objective, Research Gap, Methodology, Graph/Figure, Conclusion), and explicitly speaks to the "
        "rulebook's guidelines: clear objectives, SDG relevance, innovation, visual aids, logical flow, sound "
        "methodology, and clearly presented findings. Read it at a natural, confident pace — roughly 140–150 "
        "words per minute. Practice with a timer; if you are running long, trim the italicized 'optional detail' "
        "lines first.",
    )

    h1(doc, "Timing Plan (5 minutes total)")
    make_table(
        doc,
        ["Segment", "Target time", "Poster section it maps to"],
        [
            ["Hook + introduction", "0:00 – 0:30", "Title"],
            ["Problem & literature gap", "0:30 – 1:15", "Literature Review + Research Gap"],
            ["Research objective", "1:15 – 1:45", "Research Objective"],
            ["Methodology walk-through", "1:45 – 3:15", "Methodology"],
            ["Results & figures", "3:15 – 4:15", "Graph/Figure"],
            ["Conclusion & impact", "4:15 – 4:50", "Conclusion"],
            ["Closing line + invite questions", "4:50 – 5:00", "—"],
        ],
    )

    h1(doc, "Full Speech Script")

    h2(doc, "1. Hook and Introduction (≈30 seconds)")
    para(
        doc,
        "“Good [morning/afternoon], judges. My name is [speaker name], and on behalf of Team MatriWatch — "
        "myself, Arnob Aich Anurag, and Sadia Sultana — I'm presenting MatriWatch: Smart Maternal Health "
        "Monitoring Software for Mothers, with a specific focus today on our postpartum depression screening "
        "model.”",
    )
    para(
        doc,
        "“Every year, thousands of new mothers in Bangladesh experience postpartum depression that goes "
        "undetected — not because we lack tools like the PHQ-9 and EPDS questionnaires, but because we lack fast, "
        "data-driven ways to act on them in low-resource clinical settings. That's the gap we built MatriWatch "
        "to close.”",
    )

    h2(doc, "2. Literature Review and Research Gap (≈45 seconds)")
    para(
        doc,
        "“Postpartum depression affects not just mothers, but infant development and entire family "
        "well-being. The PHQ-9 and EPDS are the two most trusted screening instruments worldwide, and machine "
        "learning has been shown to make screening faster and more data-driven. But here's the gap: almost all "
        "existing studies either do simple binary detection — depressed or not — or they just re-predict the raw "
        "questionnaire score. Very little prior work attempts fine-grained, multi-class severity prediction for "
        "both PHQ-9 and EPDS together, on a Bangladesh-specific cohort, using explainable machine learning. "
        "That's exactly the gap MatriWatch is built to fill.”",
    )

    h2(doc, "3. Research Objective (≈30 seconds)")
    para(
        doc,
        "“Our objective was fivefold: build a Bangladesh-focused postpartum depression screening model; "
        "predict PHQ-9 severity across six fine-grained classes and EPDS risk across three levels; compare an "
        "attention-based 1D-CNN against a tuned RBF-SVM; properly handle missing data and class imbalance; and "
        "finally, explain every prediction using SHAP and feature importance — because a screening tool that "
        "clinicians can't trust or understand is a screening tool they won't use.”",
    )

    h2(doc, "4. Methodology (≈90 seconds)")
    para(
        doc,
        "“We worked with a dataset of 800 postpartum women in Bangladesh, answering 46 structured questions "
        "spanning demographics, relationships, support systems, obstetric history, and the newborn's condition. "
        "Critically, we made our model score-free — we deliberately removed the raw PHQ-9 and EPDS numeric "
        "totals from the input features, so the model has to learn from real life circumstances, not from "
        "questionnaire arithmetic.”",
    )
    para(
        doc,
        "“After cleaning, mode-imputing missing values, encoding categorical answers, and standardizing "
        "everything numerically, we balanced the training data with SMOTE and split it 70-15-15 into train, "
        "validation, and test sets.”",
    )
    para(
        doc,
        "“We then trained two competing branches in parallel. The first is an attention-based 1D-CNN: three "
        "convolution blocks that scan across the 46 features, followed by a squeeze-and-excitation attention "
        "gate that learns which feature channels matter most for each patient, and a four-layer fully-connected "
        "classifier head. The second branch is a tuned RBF Support Vector Machine. Our key methodological "
        "contribution is what we call task-adaptive selection: for each target — PHQ-9 and EPDS separately — we "
        "keep whichever branch scores higher on validation weighted F1. The model architecture isn't fixed in "
        "advance; it's chosen by evidence, per task.”",
    )
    para(
        doc,
        "*[Optional detail if time allows]: “Our final attention-CNN branch has roughly 131,000 trainable "
        "parameters per task, three convolution blocks growing from 32 to 64 to 128 channels, and the SVM "
        "branches use a tuned RBF kernel with hyperparameters re-optimized separately for PHQ-9 and EPDS.”*",
    )

    h2(doc, "5. Results and Figures (≈60 seconds)")
    para(
        doc,
        "“On the held-out test set, our final score-free model reached 80.4% accuracy and 81.6% weighted "
        "F1 on the six-class PHQ-9 task, and 82.9% accuracy and 83.0% weighted F1 on the three-class EPDS task "
        "— both substantially ahead of classical baselines like logistic regression, random forest, and "
        "gradient boosting evaluated under the exact same score-free conditions. Our confusion matrices show a "
        "strongly diagonal pattern: when the model is wrong, it's almost always confusing clinically adjacent "
        "severity levels — like Moderate versus Moderately Severe — never collapsing into a single majority "
        "guess.”",
    )
    para(
        doc,
        "“And because explainability mattered to us as much as accuracy, our SHAP analysis shows the "
        "model's top drivers are exactly the variables a clinician would expect to matter: relationship with the "
        "newborn, relationship with the husband, anger after childbirth, relationship with the in-laws, and "
        "experience of abuse. This isn't a black box — it's a model that reasons the way a clinician reasons.”",
    )

    h2(doc, "6. Conclusion and Real-World Relevance (≈35 seconds)")
    para(
        doc,
        "“In conclusion, MatriWatch shows that a score-free, task-adaptive hybrid model can deliver strong, "
        "interpretable, fine-grained postpartum depression screening — directly aligned with UN Sustainable "
        "Development Goal 3, Good Health and Well-Being, and Goal 5, Gender Equality, by giving low-resource "
        "maternal healthcare settings in Bangladesh a faster, explainable way to flag at-risk mothers before a "
        "crisis, not after one.”",
    )

    h2(doc, "7. Closing Line (≈10 seconds)")
    para(
        doc,
        "“That's MatriWatch — turning a 46-question conversation into an early, explainable warning system "
        "for the mothers who need it most. Thank you — we'd love to take your questions.”",
    )

    h1(doc, "Anticipated Judge Questions and Suggested Answers")
    qa = [
        (
            "Why remove the raw PHQ-9/EPDS score from the inputs — doesn't that throw away useful signal?",
            "Because using the raw score as an input would let the model 'cheat' by just re-deriving the label "
            "from the score's own arithmetic — that's the binary/score-shortcut problem prior work falls into. "
            "Removing it forces the model to learn from the actual life-context variables, which is far more "
            "useful for screening someone who hasn't taken the questionnaire yet, or for explaining *why* a "
            "mother is at risk, not just *that* she scored a certain number.",
        ),
        (
            "Why two different model families (CNN and SVM) instead of one?",
            "Because with about 800 rows, a deep CNN and a classical SVM have different bias-variance "
            "trade-offs, and neither one dominates on every task. Our task-adaptive selector lets the evidence — "
            "validation weighted F1 — decide per task, instead of us assuming one architecture is always best.",
        ),
        (
            "How do you handle class imbalance?",
            "SMOTE oversampling on the training split only, combined with stratified splitting so validation and "
            "test sets keep real-world class proportions, and weighted F1 as our primary evaluation metric instead "
            "of raw accuracy.",
        ),
        (
            "Is this clinically validated / ready for deployment?",
            "This is a research prototype validated on a held-out test split from our own collected cohort. The "
            "next step toward deployment is a small prospective clinical validation study and Bangla-language "
            "localization, paired with the rule-based safety thresholds already built into the broader MatriWatch "
            "platform as a fallback.",
        ),
        (
            "What is the biggest limitation of this work?",
            "Sample size — 800 respondents is workable for this kind of structured-survey classification, but a "
            "larger, multi-site Bangladesh cohort would let us validate the fine-grained 6-class PHQ-9 boundaries "
            "with more statistical confidence, especially for the rarer severity classes.",
        ),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        r = p.add_run("Q: ")
        r.bold = True
        p.add_run(q)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("A: ")
        r2.bold = True
        p2.add_run(a)
        doc.add_paragraph()

    h1(doc, "Delivery Tips Aligned to the Judging Criteria")
    bullets(
        doc,
        [
            "Clarity & Organization: follow the script's section order exactly — it mirrors the poster's own panel layout, so judges can follow your speech while glancing at the poster.",
            "Research Quality & Innovation: emphasize the score-free design and the task-adaptive selector by name — these are your two most defensible, original contributions.",
            "Visual Appeal: physically point at the relevant poster panel (confusion matrix, SHAP plot, pipeline diagram) at the moment you mention it in the script.",
            "Presentation Skills: practice against a 5-minute timer at least three times; pause briefly after each section heading rather than rushing through transitions.",
            "Industry Relevance: close with the SDG 3 / SDG 5 framing and the low-resource Bangladesh deployment angle — judges are explicitly told to weigh real-world relevance.",
        ],
    )

    out_path = OUT_DIR / "MatriWatch_Presentation_Speech.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    p1 = build_architecture_doc()
    p2 = build_speech_doc()
    print("Saved:", p1)
    print("Saved:", p2)
