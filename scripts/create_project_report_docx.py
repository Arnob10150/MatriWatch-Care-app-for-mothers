from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "MatriWatch_BEAR_Summit_Project_Report.docx"


def add_title(document: Document, title: str, subtitle: str) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(22)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_para(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if index < len(lines) - 1:
            run.add_break()


def add_metrics_table(document: Document) -> None:
    rows = [
        ("Maternal Risk Classifier", "91.61%", "91.64%", "441", "Primary Low/Mid/High maternal risk model."),
        ("Birth Weight Logistic Regression", "92.50%", "92.66%", "40", "Auxiliary pregnancy outcome model."),
        ("High-Risk Pregnancy Random Forest", "93.50%", "93.40%", "200", "High-risk pregnancy classifier."),
        ("Fetal Health Random Forest", "93.19%", "92.88%", "426", "CTG-based fetal health classifier."),
        ("Postpartum Depression Random Forest", "97.01%", "97.00%", "301", "PPD prediction support model."),
        ("GDM Ensemble Gradient Boosting", "94.27%", "94.28%", "908", "Gestational diabetes detector."),
        ("Symptom Disease TF-IDF Logistic Regression", "95.42%", "95.44%", "240", "Symptom-to-condition classifier."),
        ("Symptom Severity Triage", "100.00%*", "100.00%*", "200", "Derived clinical danger-sign target; should be framed as rule-based triage, not validated severity ML."),
    ]
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Model", "Accuracy", "Weighted F1", "Test Rows", "Use"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_report() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    add_title(
        document,
        "MatriWatch Project Report",
        "AI-powered maternal health monitoring platform for mothers, clinics, and community health workers",
    )

    add_heading(document, "1. Executive Summary")
    add_para(
        document,
        "MatriWatch is a hybrid maternal health monitoring platform designed for low-resource settings such as Bangladesh. "
        "It combines a React Native Expo mobile app for pregnant and postpartum mothers, a Next.js clinic dashboard for health workers, "
        "a Supabase backend, and a FastAPI machine learning service. Mothers submit daily vitals, symptoms, and postpartum depression "
        "screening responses. The system converts those inputs into Low, Mid, or High risk classifications and shows actionable alerts "
        "to clinics."
    )
    add_para(
        document,
        "The project is built around one practical mission: detect maternal complications before they become emergencies. "
        "The demo-ready flow shows a mother submitting a high-risk check-in, the ML service scoring the case, and the clinic dashboard "
        "surfacing the alert for staff response."
    )

    add_heading(document, "2. Problem Being Solved")
    add_para(
        document,
        "Pregnant and postpartum mothers in Bangladesh often have limited monitoring between clinic visits. Dangerous conditions such as "
        "preeclampsia, gestational diabetes, infection, fetal distress, and postpartum depression can develop silently. Paper records, "
        "phone follow-ups, and infrequent home visits are reactive and do not give clinics real-time visibility into daily patient status."
    )
    add_bullets(
        document,
        [
            "Primary users: pregnant and postpartum mothers using Android/iOS phones.",
            "Clinical users: doctors, nurses, clinic staff, and community health workers.",
            "Core value: convert daily home check-ins into clinic-visible risk alerts.",
            "Impact target: earlier detection, faster response, and fewer preventable emergencies.",
        ],
    )

    add_heading(document, "3. Product Scope")
    add_heading(document, "Mother-Facing Mobile App", 2)
    add_bullets(
        document,
        [
            "Daily vitals check-in: blood pressure, blood sugar, body temperature, heart rate, symptoms, and notes.",
            "Immediate risk result screen powered by ML service or safe local fallback rules.",
            "EPDS postpartum depression questionnaire and weekly screening support.",
            "Recovery timeline foundation for postpartum tracking.",
            "Supabase-ready authentication and data submission.",
            "Expo-based cross-platform support for Android, iOS, and mobile web development.",
        ],
    )
    add_heading(document, "Clinic-Facing Web Dashboard", 2)
    add_bullets(
        document,
        [
            "Patient list with Low, Mid, and High risk badges.",
            "Alert queue for high-risk check-ins and PPD flags.",
            "Patient detail pages with vitals history and trend charts.",
            "Population-level dashboard metrics for clinic monitoring.",
            "Supabase-ready realtime architecture for live updates.",
            "UI adapted from the supplied MatriWatch initial design package.",
        ],
    )

    add_heading(document, "4. Repository Architecture")
    add_code_block(
        document,
        [
            "MatriWatch/",
            "  apps/web          Next.js clinic dashboard",
            "  apps/mobile       Expo React Native mother app",
            "  services/ml       FastAPI ML scoring service",
            "  packages/shared   Shared TypeScript rules, types, and mock data",
            "  supabase          PostgreSQL schema, RLS policies, and Edge Function",
            "  datasets          Zipped datasets and dataset manifest",
            "  scripts           Dataset download and model training scripts",
            "  docs              Reports and project documentation",
        ],
    )

    add_heading(document, "5. Runtime Architecture")
    add_code_block(
        document,
        [
            "Mother app submits check-in",
            "-> Supabase checkins table stores vitals",
            "-> Supabase Edge Function calls FastAPI ML service",
            "-> ML service returns risk score, risk level, and reasons",
            "-> Supabase updates the check-in and creates alert rows",
            "-> Clinic dashboard reads alerts and patient status",
            "-> Clinic staff contact or escalate the mother",
        ],
    )
    add_para(
        document,
        "For local development, the system also works without Supabase credentials. The web dashboard uses demo clinical data, and the "
        "web API/mobile app can call the local FastAPI ML service directly. This makes the project suitable for live demonstrations even "
        "before production cloud credentials are configured."
    )

    add_heading(document, "6. Machine Learning Architecture")
    add_para(
        document,
        "The ML layer is implemented as a FastAPI microservice in services/ml. Training is handled by scripts/train_all_models.py, which reads "
        "the local zipped datasets and writes joblib model artifacts plus evaluation metrics into services/ml/models. The service exposes "
        "separate scoring endpoints for maternal risk, EPDS, GDM, fetal health, and symptom triage."
    )
    add_bullets(
        document,
        [
            "Primary maternal risk model: predicts Low, Mid, or High risk from vitals and demographics.",
            "PPD model: supports postpartum depression flagging alongside EPDS score rules.",
            "GDM model: classifies gestational diabetes risk using glucose, BMI, insulin, and age-style features.",
            "Fetal health model: classifies fetal status from CTG-style features.",
            "Symptom triage: combines disease classification with safety-oriented danger-sign rules.",
            "Clinical safety layer: rule thresholds remain active so obvious danger signs are not hidden by model uncertainty.",
        ],
    )

    add_heading(document, "7. Model Evaluation Metrics")
    add_metrics_table(document)
    add_para(
        document,
        "*Important caveat: the symptom severity model originally performed poorly on the raw dataset severity label, around 30.50% accuracy "
        "and 29.58% weighted F1. The current 100% score comes from a derived danger-sign target and is therefore best described as "
        "rule-based clinical triage support, not as an independently validated severity classifier."
    )

    add_heading(document, "8. Run Commands")
    add_heading(document, "Install Dependencies", 2)
    add_code_block(
        document,
        [
            "cd F:\\Tencon\\postpertum\\MatriWatch",
            "npm install",
            "python -m pip install -r services\\ml\\requirements.txt",
        ],
    )
    add_heading(document, "Train Models", 2)
    add_code_block(document, ["npm run train:models"])
    add_heading(document, "Run ML Service", 2)
    add_code_block(
        document,
        ["python -m uvicorn app.main:app --reload --app-dir services/ml --host 127.0.0.1 --port 8000"],
    )
    add_heading(document, "Run Website Version", 2)
    add_code_block(
        document,
        [
            "$env:NEXT_PUBLIC_MATRIWATCH_ML_URL=\"http://127.0.0.1:8000\"",
            "npm run dev:web",
            "Open http://localhost:3000",
        ],
    )
    add_heading(document, "Run Phone Version", 2)
    add_code_block(
        document,
        [
            "$env:EXPO_PUBLIC_MATRIWATCH_ML_URL=\"http://127.0.0.1:8000\"",
            "$env:EXPO_PUBLIC_MATRIWATCH_API_URL=\"http://localhost:3000/api\"",
            "npm run dev:mobile",
            "Then press a for Android, i for iOS, or scan the Expo QR code with Expo Go.",
        ],
    )
    add_heading(document, "Optional Direct Platform Commands", 2)
    add_code_block(
        document,
        [
            "npm --workspace apps/mobile run android",
            "npm --workspace apps/mobile run ios",
            "npm --workspace apps/mobile run web",
        ],
    )
    add_heading(document, "Verification Commands", 2)
    add_code_block(
        document,
        [
            "python -m pytest services/ml/tests",
            "npm run typecheck",
            "npm run build:web",
            "Invoke-RestMethod http://127.0.0.1:8000/health",
        ],
    )

    add_heading(document, "9. Current Implementation Status")
    add_bullets(
        document,
        [
            "The web dashboard builds successfully.",
            "The TypeScript workspace passes typechecking.",
            "The ML service tests pass.",
            "The local FastAPI service loads trained model artifacts.",
            "The web API can call the ML service and return risk results.",
            "The mobile app has check-in and EPDS flows wired to local/API scoring paths.",
            "Supabase schema and RLS policies are present, but real project credentials are required for production deployment.",
            "Push notifications and clinical pilot validation remain future hardening tasks.",
        ],
    )

    add_heading(document, "10. BEAR Summit Category A Suitability")
    add_para(
        document,
        "Public event coverage describes BEAR as a Biotechnology, Electronics, AI, and Robotics summit focused on innovation in Bangladesh. "
        "MatriWatch is well aligned with an AI/deep-tech health innovation category because AI is central to the project. If AI is removed, "
        "the product becomes a simple logging system. With AI, it becomes a proactive monitoring and alert platform."
    )
    add_bullets(
        document,
        [
            "AI centrality: risk classification, PPD prediction support, GDM detection, fetal health classification, and symptom triage.",
            "Bangladesh relevance: addresses maternal monitoring gaps in local low-resource clinical settings.",
            "Social impact: targets preventable maternal emergencies and postpartum mental health under-detection.",
            "Technical completeness: includes mobile app, web dashboard, database schema, Edge Function, and ML microservice.",
            "Demo quality: can show a high-risk case moving from phone submission to clinic alert in under 60 seconds.",
            "Scalability: designed for Supabase, Vercel, Expo, and a containerized FastAPI service.",
        ],
    )
    add_para(
        document,
        "Prototype-stage suitability score: 8.2 out of 10. The score could rise above 9 with Bangla localization, a deployed Supabase Realtime "
        "demo, a small clinical validation plan, model explainability on the dashboard, and production push notifications."
    )

    add_heading(document, "11. Recommended BEAR Demo Script")
    add_bullets(
        document,
        [
            "Open the clinic dashboard and show the patient list.",
            "Open the mother app and submit a high-risk check-in: BP 148/96, headache, swelling, 34 weeks pregnant.",
            "Show the ML service returning High Risk with reasons related to preeclampsia danger signs.",
            "Return to the dashboard and show the high-risk alert and patient detail page.",
            "Explain that the system shortens the time between home symptom reporting and clinic response.",
        ],
    )

    add_heading(document, "12. Sources")
    add_bullets(
        document,
        [
            "AIUB at the Inaugural BEAR Summit: https://www.aiub.edu/aiub-at-the-inaugural-bear-summit",
            "BRAC University Showcases Innovation at BEAR Summit 2025: https://engineering.bracu.ac.bd/news-details/brac-university-showcases-innovation-at-bear-summit-2025",
        ],
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
